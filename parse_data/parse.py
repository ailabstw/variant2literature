"""parse papers and supplementaries
"""
import os
import logging
import tempfile
import subprocess
import traceback
import zipfile
from typing import List, Dict, Any, NamedTuple

from nltk.tokenize.punkt import PunktSentenceTokenizer, PunktParameters
from bs4 import BeautifulSoup
import lxml  # pylint: disable=unused-import
import magic
import docx
import pyexcel_xls

from .utils import clean_text, timeout
from .pdf_utils import get_pdf_objects

logger = logging.getLogger(__name__)


class PaperData(NamedTuple):
    """paper data
    """
    body: str = ''
    tables: List[Dict[str, Any]] = []


@timeout(180)
def read_doc(path):
    """read .doc
    """
    with tempfile.TemporaryDirectory() as tempdir:
        with open(os.devnull, 'w') as fnull:
            subprocess.call(['soffice', '--headless', '--convert-to', 'docx',
                             '--outdir', tempdir, path], stdout=fnull, stderr=fnull)
        docx_path = os.path.join(tempdir, os.path.basename(path).rsplit('.', 1)[0] + '.docx')
        ret = read_docx(docx_path)
        os.unlink(docx_path)
    return ret


@timeout(180)
def read_docx(path):
    """read .docx (Microsoft 2007+)
    """
    try:
        doc = docx.Document(path)

        punkt_param = PunktParameters()
        punkt_param.abbrev_types = set(['fig'])
        tokenizer = PunktSentenceTokenizer(punkt_param)

        body = []
        for p in doc.paragraphs:
            body += tokenizer.tokenize(clean_text(p.text))
        body = '\n'.join(body)

        tables = []
        for t in doc.tables:
            table = {'cells': []}
            for row in t.rows:
                row_elements = []
                for cell in row.cells:
                    for p in cell.paragraphs:
                        row_elements.append({
                            'text': clean_text(p.text)
                        })
                table['cells'].append(row_elements)
            tables.append(table)

        data = PaperData(body, tables)
    except Exception:
        logger.info('fail: %s', path)
        traceback.print_exc()
        return PaperData()

    return data


@timeout(300)
def read_pdf(path, table_detect=True):
    """read pdf
    """
    try:
        body, tables = get_pdf_objects(path, table_detect)
        body = '\n'.join(body)
        data = PaperData(body, tables)

    except Exception:
        logger.info('fail: %s', path)
        traceback.print_exc()
        return PaperData()

    return data


@timeout(180)
def read_excel(path):
    """read csv, xls, xlsx
    """
    try:
        d = pyexcel_xls.get_data(path)

        tables = []
        for _, t in d.items():
            table = {'cells': []}
            max_row_len = 0
            for row in t:
                row_elements = []
                for col in row:
                    if not isinstance(col, str):
                        col = ''
                    row_elements.append({'text': clean_text(col)})
                table['cells'].append(row_elements)
                max_row_len = max(max_row_len, len(row_elements))

            for row in table['cells']:
                while len(row) < max_row_len:
                    row.append({'text': ''})
            tables.append(table)
        body = ''

        data = PaperData(body, tables)
    except Exception:
        logger.info('fail: %s', path)
        traceback.print_exc()
        return PaperData()

    return data


@timeout(180)
def read_xml(path):  # pylint: disable=too-many-locals
    """read nxml, xml, html
    """
    try:
        with open(path, 'rb') as f:
            s = f.read()
        s = s.decode('utf8')
        s = s.replace('<break/>', ', ')
        soup = BeautifulSoup(s, 'lxml')

        title = soup.find('article-title')
        title = title.getText(' ') if title is not None else ''
        title = clean_text(title)

        body = [title]
        tables = []

        punkt_param = PunktParameters()
        punkt_param.abbrev_types = set(['fig'])
        tokenizer = PunktSentenceTokenizer(punkt_param)

        for tb in soup.findAll('table'):
            table = {'cells': []}
            for tr in tb.findAll(['tr']):
                row_elements = []
                for td in tr.findAll(['td', 'th']):
                    row_elements.append({
                        'text': clean_text(td.getText(' '))
                    })
                table['cells'].append(row_elements)

            parent = tb
            while parent is not None and parent.find('label') is None:
                parent = parent.find_parent()
            if parent is not None:
                label = parent.find('label').getText(' ')
                caption_obj = parent.find('caption')
                if caption_obj is not None:
                    caption = caption_obj.getText(' ')
                else:
                    caption = ''
            else:
                label, caption = None, None

            table.update({
                'caption': {
                    'text': caption,
                    'label': label,
                }
            })
            tables.append(table)

        for paragraph in soup.findAll('p'):
            for t in paragraph.findAll('table'):
                t.extract()
            p = map(clean_text, paragraph.getText(' ').split())
            p = ' '.join(filter(bool, p))
            body += tokenizer.tokenize(p)
        body = '\n'.join(body)

        data = PaperData(body, tables)

    except Exception:
        logger.info('fail: %s', path)
        traceback.print_exc()
        return PaperData()

    return data


def read_zip(file_path, table_detect):
    """read zip file
    """
    zip_filename = os.path.basename(file_path)
    data_list = []

    with zipfile.ZipFile(file_path, 'r') as zipf:
        for name in zipf.namelist():
            if name.endswith('/'):
                continue
            ext = name.rsplit('.', 1)[-1].lower()
            with zipf.open(name) as f, tempfile.NamedTemporaryFile(suffix=f'.{ext}') as ftmp:
                ftmp.write(f.read())

                for _name, data in parse_file(ftmp.name, table_detect):
                    if ext == 'zip':
                        _name = os.path.normpath(_name)
                        name = os.path.join(name, *_name.split(os.sep)[1:])
                    filename = os.path.join(zip_filename, name)
                    data_list.append((filename, data))
    return data_list

def read_txt(file_path):
    """read txt file
    """
    try:
        with open(file_path, 'rb') as f:
            s = f.read()
        s = s.decode('utf8')
        data = PaperData(s, [])

    except Exception:
        logger.info('fail: %s', path)
        traceback.print_exc()
        return PaperData()

    return data


def parse_file(file_path, table_detect):
    """parse files in different format
    """
    filename = os.path.basename(file_path)
    ext = filename.rsplit('.', 1)[-1].lower()
    ftype = magic.from_file(file_path)

    if ext in ['pdf'] and ftype.find('PDF') >= 0:
        yield (filename, read_pdf(file_path, table_detect))

    elif ext in ['doc', 'docx'] and ftype.find('Microsoft Word 2007+') >= 0:
        yield (filename, read_docx(file_path))

    elif ext in ['doc', 'docx'] and ftype.find('Composite Document File') >= 0:
        yield (filename, read_doc(file_path))

    elif ext in ['html', 'xml', 'nxml']:
        yield (filename, read_xml(file_path))

    elif ext in ['xlsx', 'xls', 'cvs']:
        yield (filename, read_excel(file_path))

    elif ext in ['txt']:
        yield (filename, read_txt(file_path))

    elif ext in ['zip'] and ftype.find('Zip') >= 0:
        yield from read_zip(file_path, table_detect)


def parse_dir(dirname, nxml_only=False, table_detect=True):
    """parse the directory of a paper
    """
    idx = 0
    for filename in os.listdir(dirname):
        if nxml_only and not filename.endswith('.nxml'):
            continue

        path = os.path.join(dirname, filename)
        try:
            for _filename, data in parse_file(path, table_detect):
                if data:
                    yield (idx, _filename, data)
                    idx += 1
        except TimeoutError:
            logger.info(f'timeout {path}')
